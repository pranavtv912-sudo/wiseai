"""
Resume Parser Service
Extracts text and structured data from PDF and DOCX resume files
"""

import PyPDF2
from docx import Document
import re
from typing import Dict, List, Tuple
import spacy


class ResumeParser:
    """Service for parsing resume files and extracting structured data"""
    
    def __init__(self):
        """Initialize spacy model for NLP"""
        try:
            self.nlp = spacy.load('en_core_web_sm')
        except:
            print('Warning: SpaCy model not loaded. Install with: python -m spacy download en_core_web_sm')
            self.nlp = None
    
    def parse_resume(self, file_path: str) -> Dict:
        """
        Parse resume file and extract data
        
        Args:
            file_path: Path to resume file (.pdf or .docx)
            
        Returns:
            Dictionary with extracted data
        """
        try:
            if file_path.endswith('.pdf'):
                text = self._extract_text_from_pdf(file_path)
            elif file_path.endswith(('.docx', '.doc')):
                text = self._extract_text_from_docx(file_path)
            else:
                return {'error': 'Unsupported file format'}
            
            # Extract structured data from text
            extracted_data = {
                'raw_text': text,
                'skills': self._extract_skills(text),
                'experience': self._extract_experience(text),
                'education': self._extract_education(text),
                'certifications': self._extract_certifications(text),
                'projects': self._extract_projects(text),
                'contact_info': self._extract_contact_info(text),
                'keywords': self._extract_keywords(text),
            }
            
            return extracted_data
        
        except Exception as e:
            return {'error': f'Error parsing resume: {str(e)}'}
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ''
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text()
        except Exception as e:
            print(f'Error extracting PDF: {str(e)}')
        
        return text
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ''
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + '\n'
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + '\n'
        except Exception as e:
            print(f'Error extracting DOCX: {str(e)}')
        
        return text
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from resume text"""
        # Common skill keywords
        skill_keywords = {
            'Python', 'Java', 'C++', 'JavaScript', 'TypeScript', 'C#', 'PHP', 'Ruby', 'Go',
            'SQL', 'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'Firebase',
            'Flask', 'Django', 'FastAPI', 'Node.js', 'Express', 'React', 'Vue', 'Angular',
            'Docker', 'Kubernetes', 'AWS', 'Azure', 'GCP', 'Git', 'Jenkins', 'CI/CD',
            'REST API', 'GraphQL', 'Microservices', 'Agile', 'Scrum',
            'Machine Learning', 'Deep Learning', 'TensorFlow', 'PyTorch', 'Scikit-learn',
            'Data Analysis', 'Pandas', 'NumPy', 'Matplotlib', 'Seaborn',
            'HTML', 'CSS', 'Bootstrap', 'Tailwind', 'SASS',
            'Linux', 'Windows', 'macOS', 'Bash', 'PowerShell',
            'Communication', 'Leadership', 'Problem Solving', 'Team Work', 'Creativity'
        }
        
        text_upper = text.upper()
        extracted_skills = []
        
        for skill in skill_keywords:
            if skill.upper() in text_upper:
                extracted_skills.append(skill)
        
        # Use spacy for additional NLP-based extraction if available
        if self.nlp:
            doc = self.nlp(text.lower()[:5000])  # Limit text for performance
            # Additional extraction logic could go here
        
        return list(set(extracted_skills))
    
    def _extract_experience(self, text: str) -> List[Dict]:
        """Extract work experience from resume"""
        experience_list = []
        
        # Pattern to find experience sections
        patterns = [
            r'(?:Experience|Work Experience|Professional Experience)(.*?)(?:Education|Skills|Projects|Certifications|$)',
            r'(?:Employment History)(.*?)(?:Education|Skills|Projects|Certifications|$)'
        ]
        
        experience_text = ''
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
            if matches:
                experience_text = matches[0]
                break
        
        # Extract individual jobs
        job_pattern = r'([A-Z][^A-Z\n]*?(?:engineer|developer|analyst|manager|specialist)[^A-Z\n]*?)\n(.*?)(?=\n[A-Z][^A-Z\n]*?(?:engineer|developer|analyst|manager|specialist)|$)'
        jobs = re.findall(job_pattern, experience_text, re.IGNORECASE)
        
        for job_title, job_details in jobs:
            experience_list.append({
                'title': job_title.strip(),
                'details': job_details.strip()
            })
        
        return experience_list[:5]  # Limit to 5 most recent
    
    def _extract_education(self, text: str) -> List[Dict]:
        """Extract education information from resume"""
        education_list = []
        
        # Pattern for education section
        education_pattern = r'(?:Education|Academic)(.*?)(?:Experience|Skills|Projects|Certifications|$)'
        matches = re.findall(education_pattern, text, re.IGNORECASE | re.DOTALL)
        
        if matches:
            education_text = matches[0]
            
            # Look for degrees
            degree_patterns = [
                r'(?:B\.?Sc|B\.?A|B\.?E|M\.?Sc|M\.?A|M\.?E|MBA|PhD|Diploma)[^\n]*',
                r'(?:Bachelor|Master|Associate|PhD)[^,\n]*(?:in|of)[^\n]*'
            ]
            
            for pattern in degree_patterns:
                degrees = re.findall(pattern, education_text, re.IGNORECASE)
                for degree in degrees:
                    education_list.append({'degree': degree.strip()})
        
        return education_list[:3]
    
    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications from resume"""
        certifications = []
        
        # Common certification patterns
        cert_keywords = {
            'AWS', 'Azure', 'GCP', 'Kubernetes', 'Docker',
            'CCNA', 'CISSP', 'Security+',
            'SAP', 'Oracle', 'Salesforce',
            'PMP', 'Scrum Master', 'CSM', 'PMP',
            'Data Science', 'Machine Learning',
            'Google Cloud', 'AWS Certified'
        }
        
        text_upper = text.upper()
        for cert in cert_keywords:
            if cert.upper() in text_upper:
                certifications.append(cert)
        
        return list(set(certifications))
    
    def _extract_projects(self, text: str) -> List[Dict]:
        """Extract projects from resume"""
        projects = []
        
        # Pattern to find projects section
        project_pattern = r'(?:Projects|Portfolio)(.*?)(?:Experience|Skills|Certifications|Education|$)'
        matches = re.findall(project_pattern, text, re.IGNORECASE | re.DOTALL)
        
        if matches:
            projects_text = matches[0]
            # Extract project titles and descriptions
            project_items = re.findall(r'([^\n:]+)[:–-]?\s*([^\n]*(?:\n(?![A-Z])[^\n]*)*)', projects_text)
            
            for title, desc in project_items[:5]:
                projects.append({
                    'name': title.strip(),
                    'description': desc.strip()
                })
        
        return projects
    
    def _extract_contact_info(self, text: str) -> Dict:
        """Extract contact information"""
        contact_info = {}
        
        # Email pattern
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        emails = re.findall(email_pattern, text)
        if emails:
            contact_info['email'] = emails[0]
        
        # Phone pattern
        phone_pattern = r'(?:\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phones = re.findall(phone_pattern, text)
        if phones:
            contact_info['phone'] = phones[0]
        
        # LinkedIn pattern
        linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+'
        linkedin = re.findall(linkedin_pattern, text, re.IGNORECASE)
        if linkedin:
            contact_info['linkedin'] = linkedin[0]
        
        # GitHub pattern
        github_pattern = r'(?:https?://)?(?:www\.)?github\.com/[\w-]+'
        github = re.findall(github_pattern, text, re.IGNORECASE)
        if github:
            contact_info['github'] = github[0]
        
        return contact_info
    
    def _extract_keywords(self, text: str) -> List[str]:
        """Extract important keywords from resume"""
        keywords = []
        
        if self.nlp:
            doc = self.nlp(text[:5000])
            # Extract named entities and noun chunks
            for ent in doc.ents:
                if ent.label_ in ['PERSON', 'ORG', 'GPE', 'PRODUCT']:
                    keywords.append(ent.text)
            
            for chunk in doc.noun_chunks:
                if len(chunk.text) > 3:
                    keywords.append(chunk.text)
        
        return list(set(keywords))[:20]
